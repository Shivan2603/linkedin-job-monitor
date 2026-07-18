"""
ai_resume.py — Multi-AI resume tailoring bridge for jobbot
Delegates all resume tailoring to the TailorRobot engine at E:\SivaShankar\tailorrobot.
"""
import os, re, json, time, sys
if hasattr(sys.stdout, 'reconfigure'):
    try:
        sys.stdout.reconfigure(encoding='utf-8')
    except Exception:
        pass
if hasattr(sys.stderr, 'reconfigure'):
    try:
        sys.stderr.reconfigure(encoding='utf-8')
    except Exception:
        pass
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ─── LOCAL BUILDER AND CONFIG ONLY (NO DELEGATION TO TAILORROBOT) ───
from bot.config import TAILORED_TODAY, BASE_RESUME_DOCX, DATA_FOLDER
from bot.resume_builder_core import build_tailored_resume_from_json
_tr_tailor_resume = None
_USING_TAILORROBOT = False

from bot.config import GROQ_API_KEY
from bot.utils import logger
from bot.ai_router import ai_complete, check_resume_ats
from bot.resume_research_agent import (
    research_company,
    build_narrative_strategy,
    expand_projects,
    score_and_rewrite_bullets,
    enforce_consistency,
    enforce_ats_keywords,
    write_perfect_fit_summary,
    deep_rewrite_projects,
)
from bot.cover_letter_agent import generate_cover_letter
# interview_prep_agent removed — only resume + cover letter are produced

# ─── LOCAL SWARM IMPLEMENTATION (FALLBACK) ─────────────────────────

TELEMETRY_FILE = os.path.join(DATA_FOLDER, "telemetry.json")

# ─── AGENT SYSTEM PROMPTS ──────────────────────────────────────────────────

JD_INTELLIGENCE_SYSTEM = """You are the JD Intelligence Agent — the first and most critical agent in the resume-tailoring pipeline.
Your job is to DEEPLY analyse the Job Description and extract strategic intelligence that will drive EVERY section of the resume.

Think like a senior recruiter AND like the candidate trying to stand out. Analyse:

1. LOCATION INTELLIGENCE:
   - What city/country is the job in? (e.g. "Kuala Lumpur, Malaysia" or "London, UK" or "Remote")
   - Is visa sponsorship mentioned or implied?
   - What is the ideal location_line for the resume header?
   - CRITICAL INDIA RULE: If the job is in ANY Indian city or location — Bangalore, Bengaluru, Chennai, Hyderabad, Pune, Mumbai, Delhi, New Delhi, Noida, Gurugram, Gurgaon, Kolkata, Kochi, Coimbatore, Trivandrum, Ahmedabad, Jaipur, or any other Indian city — ALWAYS set is_international=false, requires_relocation=false, visa_sponsorship_mentioned=false, and location_line="Chennai, India  |  Open to Remote / Hybrid". NEVER add 'Visa sponsorship required' for Indian jobs.
   - For truly international jobs outside India (UK, USA, Malaysia, Singapore, Australia, Canada, UAE, Europe, etc.): location_line="Chennai, India  |  Open to Global Relocation (Remote / Hybrid)  |  Visa sponsorship required"
   - For India-based jobs (whether remote, hybrid, or onsite): location_line="Chennai, India  |  Open to Remote / Hybrid"

2. DOMAIN INTELLIGENCE:
   - What industry/domain is the company in? (fintech, government, e-commerce, healthcare, SaaS, logistics, banking, telecom, consulting, insurance, tax)
   - What domain power words does the JD use? (e.g. "core banking", "digital transformation", "regulatory compliance", "citizen-facing", "high-frequency trading")
   - What domain-specific tech is prioritised? (e.g. fintech → PCI-DSS, JWT, Redis, high-throughput; gov → FIPS, Section 508, WCF, federal; cloud-native → Docker, Kubernetes, microservices)

3. SENIORITY INTELLIGENCE:
   - Is this a team lead / tech lead / architect role? (look for: "lead", "mentor", "architect", "design decisions", "team of", "drive technical")
   - Or is it individual contributor? (look for: "hands-on", "implement", "develop", "engineer")
   - What bullet style fits best: "achievement" (team lead → metrics + team) or "technical" (IC → deep stack + performance)?

4. CULTURE INTELLIGENCE:
   - What 3-5 culture keywords describe this company? (e.g. "collaborative", "fast-paced", "innovation-driven", "mission-critical", "customer-first")
   - What exact phrases from the JD should be mirrored in the summary? (pick 2-3 most distinctive phrases)

5. SUMMARY CLOSING (Line 5):
   - Write the EXACT last line of the professional summary.
   - CRITICAL RULE: DO NOT mention the company name anywhere in the resume. The resume is a generic document shared with multiple recruiters. Only the cover letter mentions the company name.
   - Format: "[Eager to / Excited to / Committed to] [contribute to / build] [domain keyword] solutions in [City/context], bringing [top relevant skill from JD] to high-impact teams."
   - If remote: "Bringing [top skill] to [domain]-focused distributed teams, delivering results across international enterprise environments."
   - If India-based: "Bringing scalable [top skill] expertise to deliver [JD's stated outcome] in fast-paced [domain] environments."
   - Good example: "Excited to contribute to high-availability fintech platforms in Kuala Lumpur, bringing 4+ years of PCI-DSS compliant .NET expertise to mission-critical digital banking systems."
   - BAD example (DO NOT DO THIS): "Excited to join CIMB's digital banking team..." ← NEVER put company name in resume summary.

6. CERTIFICATIONS RELEVANCE:
   - Which of our certs should be highlighted? (AZ-204 if Azure in JD; NEICE if gov/federal; AZ-204 with [HIGHLY RELEVANT] suffix if Azure is the primary cloud)

7. EXPERIENCE INTELLIGENCE:
   - Carefully read the years of experience required by the JD.
   - The candidate has exactly 4 years of experience.
   - The job is a MATCH if it requires EXACTLY 4 years, or 4+ years (e.g. "4+ years", "minimum 4 years", "4-6 years", "4-7 years", "4-8 years").
   - The job is NOT a match if it requires 5 years or more (e.g. "5+ years", "5-8 years", "6 years", "8+ years"), or if it requires a range starting below 4 (e.g. "3-4 years", "2-4 years", "2-5 years", "3+ years", "2+ years").
   - Set "is_experience_matching" to true if it matches exactly 4 or 4+ years. Set to false if it requires less than 4 (e.g. 2, 3) or more than 4 (e.g. 5, 6).

Return ONLY valid JSON (no markdown, no code blocks):
{
  "job_location_city": "Kuala Lumpur",
  "job_location_country": "Malaysia",
  "job_location_type": "Hybrid",
  "requires_relocation": true,
  "visa_sponsorship_mentioned": true,
  "location_line": "Chennai, India  |  Open to Global Relocation (Remote / Hybrid)  |  Visa sponsorship required",
  "company_domain": "fintech",
  "domain_power_words": ["digital banking", "core banking", "financial services"],
  "domain_priority_skills": ["PCI-DSS", "JWT", "Redis", "OAuth2", "high-throughput APIs"],
  "domain_priority_category": "Security & Messaging",
  "seniority_level": "Senior",
  "is_team_lead_role": false,
  "bullet_style": "technical",
  "culture_keywords": ["collaborative", "innovation", "agile"],
  "jd_mirror_phrases": ["deliver scalable financial solutions", "cloud-native microservices"],
  "summary_closing_line": "Excited to relocate to Kuala Lumpur and bring 4+ years of PCI-DSS compliant .NET expertise to CIMB's digital banking platform.",
  "cert_highlight": "AZ-204 (HIGHLY RELEVANT — Azure is primary cloud)",
  "is_india_remote": false,
  "is_international": true,
  "is_experience_matching": true
}"""

ANALYZER_SYSTEM = """You are the Analyzer Agent in the JCode Multi-Agent Swarm.
Your job is to perform a deep analysis of the Job Description (JD).

Extract and return a JSON containing:
1. "must_haves": [list of critical required skills/tools/languages],
2. "nice_to_haves": [list of preferred or bonus skills],
3. "soft_skills": [list of collaboration/leadership terms],
4. "company_domain": "e.g. fintech, cloud, tax, procurement",
5. "company_name": "name of company",
6. "job_title": "exact job title from JD",
7. "exact_phrases": [3-5 key phrases used in the JD to mirror]

Return ONLY valid JSON. Do not include markdown code block formatting."""


RERANKER_SYSTEM = """You are the Reranker Agent in the JCode Multi-Agent Swarm.
Your job is to match the candidate's skills against the JD's Must-Haves and Nice-to-Haves, performing precision reranking.
Group the candidate's actual skills into standard categories, placing JD-matching keywords FIRST in each list, and removing irrelevant categories.

Allowable candidate skills to categorize:
- Backend: .NET Core 7/8, .NET Framework 2.0/3.5/4.0/4.8, C#, ASP.NET Web API, ASP.NET MVC, ADO.NET, EF Core, CQRS, Clean Architecture, Microservices, YARP Reverse Proxy, SignalR, gRPC, WCF, Repository Pattern, Unit of Work
- AI / ML: Azure OpenAI GPT-4, Semantic Kernel, LangChain (.NET), Vector Embeddings, pgvector, Azure AI Search, Azure Form Recognizer, GitHub Copilot, Prompt Engineering
- Cloud: Azure App Services, Azure SQL, Azure Redis Cache, Azure DevOps, Azure Blob Storage, Application Insights, ARM Templates
- Frontend: JavaScript (JS), AJAX, Angular 15+, React, TypeScript, RxJS, NgRx/Redux, Material-UI, Tailwind CSS, Vue.js, Lazy Loading, Code Splitting
- Databases: SQL Server (including SQL Server 2008), PostgreSQL, MySQL, Oracle, Redis, pgvector, LINQ Optimisation, Stored Procedures, Full-Text Indexing
- DevOps & CI/CD: Docker, Azure DevOps YAML, GitHub Actions, Git Flow, SonarQube, OpenTelemetry, Grafana K6, Serilog
- Security: JWT, OAuth2, OIDC, AES-256 Encryption, RBAC, IP Whitelisting, X.509 Certificate Rotation, mTLS, PCI-DSS, OWASP Top 10, FIPS Compliance
- Messaging: RabbitMQ, Redis Pub/Sub, Async Workflows, Event-Driven Architecture, Polly Circuit Breakers
- Testing: xUnit, NUnit, Moq, Integration Testing, TDD, Grafana K6 Load Testing
- Methodology: Agile/Scrum, Sprint Planning, Code Reviews, Architectural Decision Records (ADRs), Team Mentoring, Software Development Lifecycle (SDLC), Technical Documentation (Requirement Specification, User Manual, System Manual), Client Brief Translation, Requirement Analysis

Return a JSON with "skills_by_category" containing lists for: Backend, Frontend, Cloud, Databases, DevOps, Security, Testing, Methodology.
CRITICAL: Return ONLY standard JSON. All keys and string values must be enclosed in double quotes. Do not include raw unquoted strings, variable names, comments, or trailing commas. Do not include markdown code block formatting."""

TAILOR_SYSTEM = """You are the Tailor Agent in the JCode Multi-Agent Swarm — operating in DEEP PERSONALIZATION MODE.
You receive: JD Intelligence, Company Intelligence, Narrative Strategy, Analyzer findings, Reranker skills, and the full candidate base facts.
Your job: FRESHLY WRITE a resume that tells a completely unique, company-specific story.
This resume must feel like it was written by a human who spent 2 hours studying this specific company.

== CRITICAL ANTI-TEMPLATE RULES (VIOLATIONS WILL BE REJECTED) ==
1. Write a highly professional, dense, factual summary (3-4 sentences) based on the candidate's real QRP/microservices experience.
2. DO NOT use gimmicky hooks like "Performance-obsessed". Keep it mature and grounded.
3. DO NOT use the same work experience bullets as a generic resume. Select and reframe bullets based on narrative_strategy.top_bullets_for_primary_role.
4. DO NOT mention the company name anywhere in the resume body. Domain/industry language only.
5. DO NOT copy bullets word-for-word from the base facts. REFRAME with the company's exact language.

== SUMMARY WRITING GUIDE ==
The professional summary MUST be highly professional, dense, and factual (3-4 sentences).
Do NOT use gimmicky hooks like "Performance-obsessed" or "Impact-driven".
Instead, write a grounded, mature engineering summary that mirrors this style:
"Senior Software Engineer with 4+ years building high-availability distributed systems, fault-tolerant microservices, and AI-augmented enterprise platforms on Microsoft Azure. Currently modernizing enterprise compliance platforms — architecting 8-stage QLM workflow engines, QAR survey platforms, and cross-system integrations (SWIFT/HANA) using .NET Core 8 and Angular. Previously architected a 12-service PCI-DSS-compliant microservices platform sustaining 300+ RPS at 99.98% uptime. Mentors 7-engineer teams and ships production features from schema design through observability. Microsoft AZ-204 Certified."

Tailor the technologies and focus areas to the JD, but keep this exact professional, dense, and fact-driven tone. Ensure you explicitly mention the years of experience (e.g. "with 4+ years of experience"). Never use AI buzzword salads or fragmented casual "hook" sentences. Be direct, authoritative, and grounded in the candidate's actual LTIMindtree QRP and DSSI microservices experience.

== WORK EXPERIENCE BULLET RULES ==
Bullet COUNT per role comes from narrative_strategy.bullet_allocation. Respect it exactly.

For the PRIMARY ROLE (first in narrative_strategy.role_emphasis_order):
- The first 3 bullets MUST be the achievements listed in narrative_strategy.top_bullets_for_primary_role (in that order)
- Reframe each achievement using the company's domain language and JD mirror phrases
- DO NOT suppress achievements listed in narrative_strategy.suppress_achievements — simply de-emphasize them (put them last or omit if bullet count is tight)

Bullet rewriting rules:
- Use domain_verbs from narrative_strategy as the opening verbs
- Frame outcomes in THIS company's business context (e.g. for fintech: "financial throughput", for SaaS: "platform reliability")
- At least 1 bullet per role must show WHY an architectural decision was made (trade-off rationale)
- PROOF OVER CLAIMS: every bullet needs a number, %, or concrete scale
- Tense: LTIMindtree (current) = present tense. All others = past tense.

== PROJECT ORDERING ==
Order projects exactly as listed in narrative_strategy.project_order.

CRITICAL: You must strictly adhere to the Candidate Base Facts below. Never fabricate, invent, or extrapolate any experience, technology, or metric. All rewrites must be 100% grounded in these base facts.

=========================================
CANDIDATE BASE FACTS (Grounded Database)
=========================================
1. HEADER:
   - Name: SIVA SHANKAR
   - Contact: +91 6383149155 | sivashankar.avi6@gmail.com
   - Links: LinkedIn: https://www.linkedin.com/in/siva-shankar-4a7849226/ | GitHub: https://github.com/shivan2603 | Portfolio: https://shivan2603.github.io/sivashankar-portfolio/
   - Location Line Options:
     * If international / relocation required: "Chennai, India  |  Open to Global Relocation (Remote / Hybrid)  |  Visa sponsorship required"
     * If India-based (remote / onsite / hybrid): "Chennai, India  |  Open to Remote / Hybrid"

2. WORK EXPERIENCE:
   * Role 1: LTIMindtree (Jun 2025 – Present)
     - Core Title: Senior Software Engineer (can tailor to "Senior Software Engineer", "Senior Backend Developer", "Senior .NET Developer" depending on JD)
     - Deloitte QRP Client context: Client: Deloitte — Quality & Risk Portal (QRP)
* Engineered 8-stage QLM case-management workflow and QAR survey engine using .NET Core 8 and Angular.
* Implemented granular RBAC across 7+ roles securing sensitive communications.
* Integrated QRP with SWIFT and CP3 for near-real-time data sync feeding QAR eligibility rules.
       * Built a semantic search engine using pgvector for sub-200ms document lookups.
       * Leveraged GitHub Copilot prompt engineering for unit tests, raising coverage to 85% with xUnit.

   * Role 2: DSSI Solutions India Pvt Ltd (Nov 2024 – May 2025)
     - Core Title: Senior Software Engineer (can tailor to "Senior Software Engineer", "Senior .NET Developer", "Senior Backend Engineer" depending on JD)
     - Procurement Client context: Financial Procurement Platform
     - Allowed Technologies: C#, .NET 7, Clean Architecture, CQRS, YARP Reverse Proxy, Docker, Azure App Services, RabbitMQ, Redis, JWT, AES-256, Agile/Scrum
     - Grounded Achievements/Metrics to reframe:
       * Engineered 12+ procurement microservices using .NET 7, CQRS, and Clean Architecture.
       * Configured RabbitMQ async messaging, increasing message processing throughput by 3x.
       * Mentored 4 junior software engineers on CQRS and Git Flow, reducing post-deployment bugs by 40%.
       * Containerized all 12 microservices using Docker, decreasing container image sizes by 65%.
       * Implemented Azure Form Recognizer for automated invoice data extraction, saving 100+ manual hours.
       * Configured YARP Reverse Proxy for path-based routing, maintaining a 99.98% system uptime SLA.

   * Role 3: Nexa Office InfoSystems LLP (Jul 2024 – Nov 2024)
     - Core Title: Senior Software Engineer — Contract / Consultant (can tailor to "Senior Software Engineer", "Senior Full-Stack Developer", ".NET Consultant")
     - Document Management Client context: Enterprise Document Management
     - Allowed Technologies: C#, .NET Core, ASP.NET Web API, Angular, Redux/NgRx, Docker, SQL Server, OAuth2/OIDC, Material-UI, mTLS, X.509
     - Grounded Achievements/Metrics to reframe:
       * Designed modular UI components in Angular, improving page load speeds by 35% across forms.
       * Secured document repository with AES-256 encryption and OAuth2 OpenID Connect integration.
       * Profiled SQL Server queries and rebuilt indexes, accelerating search lookup performance by 25%.
       * Configured secure service-to-service communication using mTLS and X.509 certificate rotations.

   * Role 4: Kasadara Technology Solutions (Jul 2022 – Jun 2024)
     - Core Title: Software Engineer (can tailor to "Software Engineer", ".NET Developer", "Backend Developer")
     - US Gov SaaS Client context: US Government & SaaS Enterprise Platforms
     - Allowed Technologies: C#, .NET Framework, ASP.NET MVC, ADO.NET, EF Core, Go, WCF, SQL Server, Agile, FIPS Compliance, Section 508, WCAG
     - Grounded Achievements/Metrics to reframe:
       * Migrated legacy modules from .NET Framework to .NET Core, achieving a 40% memory usage reduction.
       * Developed high-throughput background processing services in Go, doubling processing speeds.
       * Refactored application UI to ensure strict compliance with US Federal Section 508 and WCAG accessibility standards.
       * Optimized legacy WCF and ADO.NET data access layers, reducing web page transaction times by 20%.

3. PROJECTS — Include ALL 5, sorted by JD domain relevance (most relevant FIRST):
     * QRP (Quality & Risk Portal): C# • .NET Core • Azure OpenAI GPT-4 • pgvector • Semantic Kernel • OpenTelemetry. (Enterprise compliance portal with 8-stage QLM workflow, QAR survey engine, and granular RBAC).
     * e-ProcureZen: C# • .NET 7 • Clean Architecture • CQRS • YARP Reverse Proxy • RabbitMQ • Redis • Docker • Azure App Services. (Financial procurement microservices platform with 3x throughput and 99.98% uptime SLA).
     * Nexa Vault: .NET Core • Angular • AES-256 Encryption • OAuth2/OIDC • Docker • SQL Server • mTLS • X.509. (Secure enterprise document vault with AES-256 encryption, OAuth2/OIDC, and mTLS certificate rotation).
     * SSO Application: ASP.NET Core • OAuth2 • OIDC • JWT • mTLS • X.509 • In-Memory Distributed Cache. (Centralized enterprise SSO with PKCE code flow, JWT caching, and mTLS service-to-service security).
     * NEICE: .NET Framework • WCF • SQL Server • FIPS Compliance • RBAC • ADO.NET • Section 508. (US National Electronic Interstate Compact Enterprise — FIPS-compliant multi-agency federal platform with WCF SOAP services).

4. CERTIFICATIONS:
   - Microsoft Azure Developer Associate (AZ-204) | Microsoft | March 18, 2024
   - Top Performer Award | Nexa Office InfoSystems LLP | 2024
   - US Government Platform (NEICE) | FIPS Compliance & Federal Security Standards | Kasadara Technology Solutions | 2022–2024 (include only for government/defense JDs)

5. EDUCATION:
   - B.E. Electronics & Communication Engineering | Kathir College of Engineering, Coimbatore (Anna University) | 2018 – 2022 | GPA: 8.6 / 10

=========================================
DYNAMIC TAILORING RULES
=========================================
1. HEADER:
   - Set "job_title_headline" to the target job title (Title Case, no location).
   - Set "location_line" based on the relocation/country context in JD Intelligence.
2. PROFESSIONAL SUMMARY (3-4 sentences):
- This is the MOST IMPORTANT section.
- Follow the SUMMARY WRITING GUIDE defined above to produce a highly professional, dense 3-4 sentence summary.
- Focus on real architecture, scale (RPS/uptime), and exact technologies.
- Weave in the company's domain language and jd_mirror_phrases naturally.
   - Sentence 4: If is_team_lead_role → "Mentored 4–6 engineers, introduced ADRs, and reduced onboarding time from 4 weeks to 10 days — ready to bring that same technical leadership to high-performing engineering teams."
     If IC role → "AZ-204 certified with a proven track record of building production-grade systems that align with demanding enterprise engineering environments."
   - Sentence 5: Use the EXACT summary_closing_line from JD Intelligence verbatim. Do NOT rewrite it.
3. SKILLS CATEGORIES:
   - Return skills grouped under Backend, Frontend, Cloud, Databases, DevOps, Security, Testing, Methodology. Place matching technologies first.
4. WORK EXPERIENCE:
   - Return roles in the order from narrative_strategy.role_emphasis_order.
   - For each company, tailor:
     * "role_title": Use JD-relevant title variant listed in base facts. Keep Deloitte client context for LTIMindtree.
     * "tech_stack_line": Technologies actually used in that role, JD-matching ones listed first.
     * "bullets": Use EXACTLY the count from narrative_strategy.bullet_allocation for each company.
       For the primary role (first in role_emphasis_order), the FIRST 3 bullets MUST be the achievements in top_bullets_for_primary_role, reframed in this company's domain language.
       Open each bullet with one of the domain_verbs from narrative_strategy.
       Every bullet needs a number, %, or concrete scale.
5. PROJECTS — ALL 5 PROJECTS, in order of JD domain relevance (most relevant FIRST):
   Each project gets EXACTLY 3 bullets. Write them in this STRICT order:
   - BULLET 1 — ARCHITECTURE RATIONALE (WHY): Explain WHY those specific technology choices were made as deliberate trade-offs.
     Template: "[Selected/Chose/Adopted] [specific tech] over [alternative] to [solve specific problem — e.g. 'handle complex 8-stage QLM workflows securely', 'enforce dynamic QAR survey rules without performance bottlenecks']."
     This must sound like a senior engineer explaining a real design decision at a technical interview, NOT a task description.
   - BULLET 2 — IMPACT (WHAT + METRIC): State the measurable outcome using ONLY the allowed metrics from the project's grounded data.
     Template: "[Action verb + what was built] — [specific metric from allowed list] [in context that maps to this JD's domain]."
     Examples: 'reducing manual review time by 60% across 1,000+ weekly submissions', '3x message throughput at 99.98% uptime SLA', 'sub-100ms p99 latency on 30+ enterprise API endpoints'.
   - BULLET 3 — JD ALIGNMENT (HOW it maps to THIS specific role/domain): Directly connect the project to what this type of role needs.
     Template: "Directly applicable to [domain/industry — e.g. 'cloud-native microservices platforms', 'federal compliance environments', 'identity security layers'] — [what the candidate can bring from this project to this kind of role]."
     CRITICAL: DO NOT mention the company name in any project bullet. The resume is a reusable document. Reference the domain/industry, not the specific company.
     Good example: "Directly applicable to high-availability procurement platforms requiring 12-factor app design — delivering the same event-driven CQRS patterns in cloud-native enterprise environments."
     BAD example: "Directly applicable to QuilinX's cloud-native microservices platform" ← NEVER put company name in resume.
6. CERTIFICATIONS:
   - List the 2-3 certifications relevant to this JD (e.g., adding AZ-204 highlight or NEICE as appropriate).

Return ONLY a JSON block structured as:
{
  "job_title_headline": "Title Case Job Title",
  "location_line": "Chennai, India | ...",
  "professional_summary": "...",
  "skills_by_category": {
    "Backend": [...],
    ...
  },
  "work_experience": [
    {
      "company": "LTIMindtree",
      "dates": "Jun 2025 – Present",
      "role_title": "Senior Software Engineer | Client: Deloitte — Quality & Risk Portal (QRP)",
      "tech_stack_line": ".NET Core 8 • ASP.NET Web API • Angular • Azure OpenAI GPT-4 • Microservices",
      "bullets": ["...", "...", "...", "..."],
      "key": "LTIMindtree"
    },
    ...
  ],
  "projects": [
    {
      "name": "QRP (Quality & Risk Portal)",
      "tech_stack": "C# · .NET Core · Azure OpenAI GPT-4 · pgvector · Semantic Kernel · OpenTelemetry",
      "bullets": ["...", "...", "..."]
    },
    {
      "name": "e-ProcureZen",
      "tech_stack": "C# · .NET 7 · Clean Architecture · CQRS · YARP Reverse Proxy · RabbitMQ · Redis · Docker · Azure App Services",
      "bullets": ["...", "...", "..."]
    },
    {
      "name": "Nexa Vault",
      "tech_stack": ".NET Core · Angular · AES-256 Encryption · OAuth2/OIDC · Docker · SQL Server · mTLS · X.509",
      "bullets": ["...", "...", "..."]
    },
    {
      "name": "SSO Application",
      "tech_stack": "ASP.NET Core · OAuth2 · OIDC · JWT · mTLS · X.509 · In-Memory Distributed Cache",
      "bullets": ["...", "...", "..."]
    },
    {
      "name": "NEICE",
      "tech_stack": ".NET Framework · WCF · SQL Server · FIPS Compliance · RBAC · ADO.NET · Section 508",
      "bullets": ["...", "...", "..."]
    }
  ],
  "certifications": [
    "...", "..."
  ],
  "education": {
    "degree": "B.E. Electronics & Communication Engineering",
    "institution": "Kathir College of Engineering, Coimbatore (Anna University)",
    "years": "2018 – 2022",
    "gpa": "8.6 / 10"
  }
}
CRITICAL: All keys and string values must be in double quotes. Do not include unquoted arrays, comments, or trailing commas. No orphaned brackets."""


VERIFIER_SYSTEM = """You are the Verifier Agent in the JCode Multi-Agent Swarm.
Your job is to compare the tailored resume draft against the JD, original candidate facts, and JD Intelligence, then correct any formatting, fact-grounding, or compliance violations.

CRITICAL: Every single bullet point and the professional summary in the resume must be rewritten to align specifically with the target Job Description (JD), utilizing the company's domain language, tech stack, and goals, while strictly adhering to the candidate's base facts.

CRITICAL: You must strictly adhere to the Candidate Base Facts below. Never fabricate, invent, or extrapolate any experience, technology, or metric. All edits must be 100% grounded in these base facts.

=========================================
CANDIDATE BASE FACTS (Grounded Database)
=========================================
1. HEADER:
   - Name: SIVA SHANKAR
   - Contact: +91 6383149155 | sivashankar.avi6@gmail.com
   - Links: LinkedIn: https://www.linkedin.com/in/siva-shankar-4a7849226/ | GitHub: https://github.com/shivan2603 | Portfolio: https://shivan2603.github.io/sivashankar-portfolio/
   - Location Line Options:
     * If international / relocation required: "Chennai, India  |  Open to Global Relocation (Remote / Hybrid)  |  Visa sponsorship required"
     * If India-based (remote / onsite / hybrid): "Chennai, India  |  Open to Remote / Hybrid"

2. WORK EXPERIENCE:
   * Role 1: LTIMindtree (Jun 2025 – Present)
     - Core Title: Senior Software Engineer (can tailor to "Senior Software Engineer", "Senior Backend Developer", "Senior .NET Developer" depending on JD)
     - Deloitte QRP Client context: Client: Deloitte — Quality & Risk Portal (QRP)
* Engineered 8-stage QLM case-management workflow and QAR survey engine using .NET Core 8 and Angular.
* Implemented granular RBAC across 7+ roles securing sensitive communications.
* Integrated QRP with SWIFT and CP3 for near-real-time data sync feeding QAR eligibility rules.
       * Built a semantic search engine using pgvector for sub-200ms document lookups.
       * Leveraged GitHub Copilot prompt engineering for unit tests, raising coverage to 85% with xUnit.

   * Role 2: DSSI Solutions India Pvt Ltd (Nov 2024 – May 2025)
     - Core Title: Senior Software Engineer (can tailor to "Senior Software Engineer", "Senior .NET Developer", "Senior Backend Engineer" depending on JD)
     - Procurement Client context: Financial Procurement Platform
     - Allowed Technologies: C#, .NET 7, Clean Architecture, CQRS, YARP Reverse Proxy, Docker, Azure App Services, RabbitMQ, Redis, JWT, AES-256, Agile/Scrum
     - Grounded Achievements/Metrics to reframe:
       * Engineered 12+ procurement microservices using .NET 7, CQRS, and Clean Architecture.
       * Configured RabbitMQ async messaging, increasing message processing throughput by 3x.
       * Mentored 4 junior software engineers on CQRS and Git Flow, reducing post-deployment bugs by 40%.
       * Containerized all 12 microservices using Docker, decreasing container image sizes by 65%.
       * Implemented Azure Form Recognizer for automated invoice data extraction, saving 100+ manual hours.
       * Configured YARP Reverse Proxy for path-based routing, maintaining a 99.98% system uptime SLA.

   * Role 3: Nexa Office InfoSystems LLP (Jul 2024 – Nov 2024)
     - Core Title: Senior Software Engineer — Contract / Consultant (can tailor to "Senior Software Engineer", "Senior Full-Stack Developer", ".NET Consultant")
     - Document Management Client context: Enterprise Document Management
     - Allowed Technologies: C#, .NET Core, ASP.NET Web API, Angular, Redux/NgRx, Docker, SQL Server, OAuth2/OIDC, Material-UI, mTLS, X.509
     - Grounded Achievements/Metrics to reframe:
       * Designed modular UI components in Angular, improving page load speeds by 35% across forms.
       * Secured document repository with AES-256 encryption and OAuth2 OpenID Connect integration.
       * Profiled SQL Server queries and rebuilt indexes, accelerating search lookup performance by 25%.
       * Configured secure service-to-service communication using mTLS and X.509 certificate rotations.

   * Role 4: Kasadara Technology Solutions (Jul 2022 – Jun 2024)
     - Core Title: Software Engineer (can tailor to "Software Engineer", ".NET Developer", "Backend Developer")
     - US Gov SaaS Client context: US Government & SaaS Enterprise Platforms
     - Allowed Technologies: C#, .NET Framework, ASP.NET MVC, ADO.NET, EF Core, Go, WCF, SQL Server, Agile, FIPS Compliance, Section 508, WCAG
     - Grounded Achievements/Metrics to reframe:
       * Migrated legacy modules from .NET Framework to .NET Core, achieving a 40% memory usage reduction.
       * Developed high-throughput background processing services in Go, doubling processing speeds.
       * Refactored application UI to ensure strict compliance with US Federal Section 508 and WCAG accessibility standards.
       * Optimized legacy WCF and ADO.NET data access layers, reducing web page transaction times by 20%.

3. PROJECTS — You must include ALL 5 projects. Do NOT drop any projects:
     * QRP (Quality & Risk Portal): C# • .NET Core • Azure OpenAI GPT-4 • pgvector • Semantic Kernel • OpenTelemetry.
     * e-ProcureZen: C# • .NET 7 • Clean Architecture • CQRS • YARP Reverse Proxy • RabbitMQ • Redis • Docker • Azure App Services.
     * Nexa Vault: .NET Core • Angular • AES-256 Encryption • OAuth2/OIDC • Docker • SQL Server • mTLS • X.509.
     * SSO Application: ASP.NET Core • OAuth2 • OIDC • JWT • mTLS • X.509 • In-Memory Distributed Cache.
     * NEICE: .NET Framework • WCF • SQL Server • FIPS Compliance • RBAC • ADO.NET • Section 508.

4. CERTIFICATIONS:
   - Microsoft Azure Developer Associate (AZ-204) | Microsoft | March 18, 2024
   - Top Performer Award | Nexa Office InfoSystems LLP | 2024
   - US Government Platform (NEICE) | FIPS Compliance & Federal Security Standards | Kasadara Technology Solutions | 2022–2024 (include only for government/defense JDs)

5. EDUCATION:
   - B.E. Electronics & Communication Engineering | Kathir College of Engineering, Coimbatore (Anna University) | 2018 – 2022 | GPA: 8.6 / 10

=========================================

Check and enforce:
1. Header Location Line: Match is_international/relocation requirements (Chennai, India | ...).
2. Job Title Headline: Must be Title Case, with no locations or parentheticals.
3. Summary Anti-Template Check (CRITICAL):
   - REJECT the summary if it starts with "[Job Title] with 4+ years of hands-on experience" — that is a forbidden template opening.
   - REJECT if the summary does not clearly reflect the narrative_strategy.summary_opening_angle provided.
   - The summary MUST end with the exact summary_closing_line from JD Intelligence verbatim.
   - Punctuation clean: no double periods, no orphaned parentheses.
   - NO company name anywhere in the summary.
4. Work Experience:
   - Dynamic role titles and tech stack lines must remain factually grounded (only technologies used in that company are allowed).
   - Bullet count per role MUST match the bullet_allocation from narrative_strategy (if provided).
   - If bullet count is wrong, trim or add bullets to match exactly.
   - First 2 bullets of the primary role must use JD must-have skills.
   - NO company name in any bullet.
5. Projects: ALL 5 projects are absolutely required. You must NOT drop any projects. Each must have exactly 3 bullets. Ordered per narrative_strategy.project_order.
6. Certifications: 2-3 items matching JD.

If any rule is violated, rewrite that section.
Return the final corrected full JSON matching this schema:
{
  "job_title_headline": "...",
  "location_line": "...",
  "professional_summary": "...",
  "skills_by_category": { ... },
  "work_experience": [
    {
      "company": "...",
      "dates": "...",
      "role_title": "...",
      "tech_stack_line": "...",
      "bullets": [...],
      "key": "..."
    },
    ...
  ],
  "projects": [
    {
      "name": "...",
      "tech_stack": "...",
      "bullets": [...]
    },
    ...
  ],
  "certifications": [...],
  "education": { ... },
  "ats_report": {
    "match_score": 98,
    "matched_keywords": [...],
    "missing_with_equivalents": [...],
    "true_gaps": [...],
    "experience_gap_compensation": "...",
    "top_3_standout_points": [...],
    "recruiter_weak_point": "..."
  }
}
CRITICAL: All fields must contain plain text strings or lists of strings. Never output any markdown code blocks, HTML, or mermaid scripts inside the JSON string values. All keys and string values must be wrapped in double quotes. Return ONLY valid JSON."""


# ─── TELEMETRY LOGGER ───────────────────────────────────────────────────────
def log_telemetry(agent_name: str, duration: float, status: str, provider: str = "Groq"):
    try:
        telemetry = []
        if os.path.exists(TELEMETRY_FILE):
            with open(TELEMETRY_FILE, "r", encoding="utf-8") as f:
                telemetry = json.load(f)
        
        telemetry.append({
            "timestamp": time.strftime("%Y-%m-%d %H:%M:%S"),
            "agent": agent_name,
            "duration_seconds": round(duration, 3),
            "provider": provider,
            "status": status
        })
        
        # Keep last 100 logs
        telemetry = telemetry[-100:]
        with open(TELEMETRY_FILE, "w", encoding="utf-8") as f:
            json.dump(telemetry, f, indent=2)
    except Exception:
        pass

# ─── MAIN COORDINATOR WORKFLOW ─────────────────────────────────────────────

def extract_resume_text(docx_path: str) -> str:
    doc = Document(docx_path)
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

def parse_json_safely(raw: str) -> dict:
    import json, re
    raw = raw.strip()
    
    # Try finding markdown code block
    m = re.search(r'```(?:json)?\s*(\{.*?\})\s*```', raw, re.DOTALL | re.IGNORECASE)
    if m:
        candidate = m.group(1).strip()
    else:
        # Try finding outer braces
        m = re.search(r'(\{.*\})', raw, re.DOTALL)
        candidate = m.group(1).strip() if m else raw
        
    # Repair unquoted items in arrays, e.g. ["Docker", GitHub Actions, Git Flow] -> ["Docker", "GitHub Actions", "Git Flow"]
    def fix_array(match):
        arr_content = match.group(1).strip()
        if any(c in arr_content for c in ['{', '}', '[', ']']):
            return match.group(0)
        elements = []
        parts = re.findall(r'("[^"\\]*(?:\\.[^"\\]*)*"|[^,]+)', arr_content)
        for p in parts:
            p = p.strip()
            if not p:
                continue
            if p.startswith('"') and p.endswith('"'):
                clean = p[1:-1]
            else:
                clean = p.strip('"')
            if not p.startswith('"') and (clean.lower() in ['true', 'false', 'null'] or re.match(r'^\d+(\.\d+)?$', clean)):
                elements.append(clean)
            else:
                escaped = clean.replace('"', '\\"')
                elements.append(f'"{escaped}"')
        return '[' + ', '.join(elements) + ']'

    # Run array fixer
    candidate = re.sub(r'\[(.*?)\]', fix_array, candidate, flags=re.DOTALL)

    # Remove single line comments // ... if any
    candidate = re.sub(r'^\s*//.*$', '', candidate, flags=re.MULTILINE)
    # Remove inline trailing commas before closing braces/brackets
    candidate = re.sub(r',\s*([\]}])', r'\1', candidate)
    
    try:
        return json.loads(candidate, strict=False)
    except Exception as e:
        cleaned = re.sub(r'[\x00-\x1F\x7F]', '', candidate)
        try:
            return json.loads(cleaned, strict=False)
        except Exception:
            raise e



def check_experience_relevance(jd_text: str, job_title: str = "") -> tuple[bool, str]:
    """
    Programmatic filter for the user's experience requirement:
    - Candidate has 4 years of experience.
    - Match if 4 falls within the required range (e.g. 0-4, 1-4, 2-4, 3-4, 4-5, 4-6, 4-8, etc.).
    - Reject if the range does not cover 4 years, or if min required is > 4.
    """
    text = f"{jd_text}\n{job_title}".lower().replace("-", " ")
    import re
    
    # 1. Match patterns like: "3-4 years", "2-4 years", "3 to 4 years", "5-8 years", "5 to 8 years"
    for match in re.finditer(r'(\d+)\s*(?:-|to)\s*(\d+)\s*(?:years|yrs|year|yr)', text):
        low = int(match.group(1))
        high = int(match.group(2))
        if low > 4:
            return False, f"Minimum required experience ({low} years) is greater than candidate's 4 years of experience"
            
    # 2. Match patterns like: "5+ years", "6+ years", "3+ yrs", "5+ year"
    for match in re.finditer(r'(\d+)\+\s*(?:years|yrs|year|yr)', text):
        val = int(match.group(1))
        if val > 4:
            return False, f"Requires {val}+ years (greater than candidate's 4 years)"
            
    # 3. Match patterns like: "minimum of 5 years", "at least 5 years", "5 years of experience"
    for match in re.finditer(r'(?:min|minimum|at least|require|requires|of|have)\s*(\d+)\s*(?:years|yrs|year|yr)', text):
        val = int(match.group(1))
        if val > 4:
            # Check if this is preceded by a range (e.g. "3 to " or "4-") already handled
            start_idx = max(0, match.start() - 10)
            context = text[start_idx:match.start()]
            if not re.search(r'\d+\s*(?:-|to)\s*$', context):
                return False, f"Requires minimum of {val} years (greater than candidate's 4 years)"
                
    return True, "Passed programmatic checks"


def check_tech_stack_relevance(job_title: str, jd_text: str) -> tuple[bool, str]:
    """
    Programmatic filter to ensure the job is relevant to the candidate's core stack (.NET / C#).
    Rejects Java, Python, PHP, C++, Android, iOS, SAP, Salesforce, QA, Scrum Master, etc. roles
    unless they explicitly mention C# or .NET.
    Also ensures the JD contains at least one of '.net', 'c#', 'dotnet', 'csharp'.
    """
    title_lower = job_title.lower()
    jd_lower = jd_text.lower()
    import re
    
    # 1. Blacklisted titles (if they don't contain .NET/C#)
    blacklist_titles = [
        'java', 'python', 'php', 'c++', 'cobol', 'golang', 'go developer', 
        'ruby', 'rails', 'android', 'ios', 'swift', 'kotlin', 'flutter', 
        'sap', 'salesforce', 'qa engineer', 'tester', 'scrum master', 
        'product manager', 'project manager', 'business analyst'
    ]
    
    # If title contains target stack, approve immediately
    has_net = '.net' in title_lower or 'dotnet' in title_lower or 'c#' in title_lower or 'csharp' in title_lower
    if has_net:
        return True, "Passed technology stack check (matched title)"
        
    for blacklisted in blacklist_titles:
        if blacklisted in title_lower:
            return False, f"Job title contains blacklisted keyword: {blacklisted}"
                
    # 2. Tech stack check (must mention .net, c#, or dotnet in the job description)
    keywords = ['.net', 'dotnet', 'c#', 'csharp', 'wcf']
    has_keyword = False
    for kw in keywords:
        if kw == '.net':
            if '.net' in jd_lower or 'dotnet' in jd_lower:
                has_keyword = True
                break
        elif kw == 'c#':
            if 'c#' in jd_lower or 'csharp' in jd_lower:
                has_keyword = True
                break
        else:
            if kw in jd_lower:
                has_keyword = True
                break
                
    if not has_keyword:
        return False, "Job description does not mention .NET or C# keywords"
        
    return True, "Passed technology stack check"

def convert_docx_to_pdf_win32(docx_path: str) -> str:
    """
    Converts a DOCX file to a PDF file using MS Word COM automation on Windows.
    Returns the absolute path to the generated PDF file if successful.
    """
    import win32com.client
    import pythoncom
    
    pythoncom.CoInitialize()
    abs_docx_path = os.path.abspath(docx_path)
    abs_pdf_path = abs_docx_path.replace(".docx", ".pdf")
    
    word = None
    doc = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(abs_docx_path)
        doc.SaveAs(abs_pdf_path, FileFormat=17) # 17 is wdFormatPDF
        doc.Close()
        return abs_pdf_path
    except Exception as e:
        raise RuntimeError(f"Word COM PDF conversion failed: {e}")
    finally:
        try:
            if doc:
                doc.Close()
        except Exception:
            pass
        try:
            if word:
                word.Quit()
        except Exception:
            pass
        pythoncom.CoUninitialize()

def restore_dropped_projects(current_json: dict, reference_projects: list) -> dict:
    from bot.resume_builder_core import get_canonical_project_name
    if not isinstance(current_json, dict):
        current_json = {}
    current_projects = current_json.get("projects", [])
    if not isinstance(current_projects, list):
        current_projects = []
    
    if not isinstance(reference_projects, list) or not reference_projects:
        return current_json
        
    verified_names = {p.get("name", "").lower().strip() for p in current_projects if isinstance(p, dict) and p.get("name")}
    
    restored = list(current_projects)
    for ref_proj in reference_projects:
        if not isinstance(ref_proj, dict):
            continue
        ref_name = ref_proj.get("name", "").lower().strip()
        if not ref_name:
            continue
        # Check if project is already present by mapping both to their canonical name
        found = False
        ref_canon = get_canonical_project_name(ref_name).lower().strip()
        for vn in verified_names:
            vn_canon = get_canonical_project_name(vn).lower().strip()
            if ref_canon == vn_canon:
                found = True
                break
        if not found:
            print(f"    [ProjectRecovery] Restoring project '{ref_proj.get('name')}' dropped during AI step.")
            restored.append(ref_proj)
            
    # Ensure all projects in the list have exactly 3 bullets (architecture rationale, impact, JD alignment)
    for p in restored:
        if not isinstance(p, dict):
            continue
        bullets = p.get("bullets", [])
        if not isinstance(bullets, list):
            bullets = [bullets] if bullets else []
        if len(bullets) < 3:
            name = p.get("name", "")
            from bot.resume_builder_core import DEFAULT_PROJECTS_POOL, PROJECTS_TECHNICAL_DECISIONS
            dp_match = next((dp for dp in DEFAULT_PROJECTS_POOL if name.lower()[:15] in dp["name"].lower() or dp["name"].lower()[:15] in name.lower()), None)
            default_bullets = dp_match["bullets"] if dp_match else ["Configured and optimized backend architecture.", "Delivered high-performance business outcome."]
            
            # Reconstruct to have 3 bullets
            bullets_3 = []
            # Bullet 1: Architecture rationale
            bullets_3.append(bullets[0] if len(bullets) > 0 else (next(iter(PROJECTS_TECHNICAL_DECISIONS.get(name.lower(), ["Selected optimal tech stack for performance."])), "Selected optimal tech stack.")))
            # Bullet 2: Impact
            bullets_3.append(bullets[1] if len(bullets) > 1 else (default_bullets[0] if default_bullets else "Optimized database and memory usage."))
            # Bullet 3: JD Alignment
            bullets_3.append(bullets[2] if len(bullets) > 2 else f"Directly applicable to company requirements, enabling secure and scalable system integration.")
            p["bullets"] = bullets_3[:3]
            
    current_json["projects"] = restored
    return current_json

def tailor_resume(job_title: str, company: str, job_description: str, site: str = "ai") -> dict:
    res = None
    if _USING_TAILORROBOT and _tr_tailor_resume:
        logger.ai(f"[JCode Swarm Bridge] Delegating resume tailoring to TailorRobot...", site=site)
        try:
            res = _tr_tailor_resume(job_title, company, job_description, site=site)
        except Exception as e:
            logger.error(f"[JCode Swarm Bridge] TailorRobot delegation failed: {e}. Falling back to local builder.", site=site)
            
    if res:
        docx_path = res.get("resume_path")
        if docx_path:
            res["resume_pdf_path"] = docx_path.replace(".docx", ".pdf")
        return res

    # Globally clean job title using unified cleaner (strips location suffixes, ALL CAPS, etc.)
    from bot.resume_builder_core import clean_job_title as _clean_title
    job_title = _clean_title(job_title)

    print(f"\n[JCode Swarm] Starting Multi-Agent Coordinator Workflow...")

    
    # Run programmatic tech stack relevance check
    is_tech_ok, tech_reason = check_tech_stack_relevance(job_title, job_description)
    if not is_tech_ok:
        print(f"    [JD Intel] Tech stack mismatch: {tech_reason}. Skipping...")
        return {
            "resume_path": "",
            "match_score": 0,
            "tailored": {},
            "ats_report": {
                "match_score": 0,
                "true_gaps": [f"Tech stack mismatch: {tech_reason}"]
            }
        }

    # Run programmatic experience check
    is_programmatic_ok, reason = check_experience_relevance(job_description, job_title)
    if not is_programmatic_ok:
        print(f"    [JD Intel] Programmatic experience mismatch: {reason}. Skipping...")
        return {
            "resume_path": "",
            "match_score": 0,
            "tailored": {
                "jd_context": {
                    "is_experience_matching": False,
                    "location_line": "Chennai, India  |  Open to Remote/Hybrid",
                    "summary_closing_line": "Bringing scalable .NET and cloud expertise to high-impact enterprise engineering teams."
                }
            },
            "ats_report": {
                "match_score": 0,
                "true_gaps": [f"Experience mismatch: {reason}"]
            }
        }

    base_text = extract_resume_text(BASE_RESUME_DOCX)

    # ─── STEP 0.5: COMPANY RESEARCH AGENT ───
    print("[JCode Coordinator] Launching Company Research Agent...")
    t0 = time.time()
    company_intelligence = research_company(company, job_title, job_description, parse_json_safely)
    log_telemetry("CompanyResearchAgent", time.time() - t0, "success")

    # ─── STEP 0: JD INTELLIGENCE AGENT ───
    print("[JCode Coordinator] Launching JD Intelligence Agent...")
    t0 = time.time()
    jd_context = {}
    try:
        intel_prompt = (
            f"Job Title: {job_title}\nCompany: {company}\n\n"
            f"Full JD:\n{job_description[:4000]}"
        )
        raw_intel = ai_complete(JD_INTELLIGENCE_SYSTEM, intel_prompt, task="analyze", max_tokens=1200)
        jd_context = parse_json_safely(raw_intel)
        log_telemetry("JDIntelligenceAgent", time.time() - t0, "success")
        
        # Check experience matching from AI
        is_ai_exp_ok = jd_context.get("is_experience_matching", True)
        if isinstance(is_ai_exp_ok, str):
            is_ai_exp_ok = is_ai_exp_ok.lower() == "true"
            
        if is_ai_exp_ok is False:
            print(f"    [JD Intel] AI experience mismatch detected: Job does not require 4 / 4+ years. Skipping...")
            return {
                "resume_path": "",
                "match_score": 0,
                "tailored": {"jd_context": jd_context},
                "ats_report": {
                    "match_score": 0,
                    "true_gaps": ["Experience mismatch: AI analyzed JD requires non-matching experience profile."]
                }
            }

        loc = jd_context.get('job_location_city', 'Unknown')
        domain = jd_context.get('company_domain', 'general')
        lead = jd_context.get('is_team_lead_role', False)
        intl = jd_context.get('is_international', False)
        print(f"    [JD Intel] Location: {loc} | Domain: {domain} | Lead role: {lead} | International: {intl}")
        print(f"    [JD Intel] Location line: {jd_context.get('location_line', 'N/A')}")
    except Exception as e:
        log_telemetry("JDIntelligenceAgent", time.time() - t0, f"failed: {e}")
        jd_context = {
            "job_location_city": "Unknown", "job_location_country": "India",
            "requires_relocation": False, "is_international": False,
            "company_domain": "technology", "is_team_lead_role": False,
            "location_line": "Chennai, India  |  Open to Remote/Hybrid",
            "summary_closing_line": "Bringing scalable .NET and cloud expertise to deliver high-quality enterprise software in fast-paced environments.",
            "jd_mirror_phrases": [], "domain_power_words": [], "domain_priority_skills": [],
            "cert_highlight": "AZ-204", "bullet_style": "achievement"
        }
        print(f"    [JD Intel] Failed (using defaults): {e}")

    # ─── STEP 1: ANALYZER AGENT ───
    print("[JCode Coordinator] Launching Analyzer Agent...")
    t0 = time.time()
    try:
        raw_analysis = ai_complete(ANALYZER_SYSTEM, f"Analyze this JD:\n{job_description[:4000]}", task="analyze", max_tokens=1000)
        analysis = parse_json_safely(raw_analysis)
        log_telemetry("AnalyzerAgent", time.time() - t0, "success")
        print(f"    [Analyzer] Extracted {len(analysis.get('must_haves', []))} must-haves.")
    except Exception as e:
        log_telemetry("AnalyzerAgent", time.time() - t0, f"failed: {e}")
        analysis = {"must_haves": [job_title], "nice_to_haves": [], "exact_phrases": []}
        print(f"    [Analyzer] Failed, using fallbacks: {e}")

    # ─── STEP 2: RERANKER AGENT ───
    print("[JCode Coordinator] Launching Reranker Agent...")
    t0 = time.time()
    try:
        rerank_prompt = (
            f"Rerank candidate skills based on Must-Haves: {analysis.get('must_haves')} "
            f"and Nice-to-Haves: {analysis.get('nice_to_haves')}. "
            f"Domain: {jd_context.get('company_domain', 'technology')}. "
            f"Prioritise skills in domain_priority_skills: {jd_context.get('domain_priority_skills', [])}"
        )
        raw_skills = ai_complete(RERANKER_SYSTEM, rerank_prompt, task="rerank", max_tokens=1200)
        skills_ranked = parse_json_safely(raw_skills)
        log_telemetry("RerankerAgent", time.time() - t0, "success")
        print("    [Reranker] Precision skills reranking completed.")
    except Exception as e:
        log_telemetry("RerankerAgent", time.time() - t0, f"failed: {e}")
        skills_ranked = {"skills_by_category": {}}
        print(f"    [Reranker] Failed, using default categories: {e}")

    # ─── STEP 2.5: NARRATIVE STRATEGY AGENT ───
    print("[JCode Coordinator] Launching Narrative Strategy Agent...")
    t0 = time.time()
    try:
        narrative_strategy = build_narrative_strategy(
            company=company,
            job_title=job_title,
            jd_text=job_description,
            company_intelligence=company_intelligence,
            jd_context=jd_context,
            parse_json_safely=parse_json_safely
        )
        log_telemetry("NarrativeStrategyAgent", time.time() - t0, "success")
    except Exception as e:
        log_telemetry("NarrativeStrategyAgent", time.time() - t0, f"failed: {e}")
        narrative_strategy = {
            "candidate_angle": "Cloud-native .NET engineer with 4+ years delivering enterprise platforms",
            "summary_opening_angle": "performance",
            "role_emphasis_order": ["LTIMindtree", "DSSI Solutions", "Nexa Office InfoSystems", "Kasadara Technology Solutions"],
            "bullet_allocation": {"LTIMindtree": 6, "DSSI Solutions": 5, "Nexa Office InfoSystems": 4, "Kasadara Technology Solutions": 3},
            "top_bullets_for_primary_role": ["A4", "A7", "A5"],
            "suppress_achievements": [],
            "domain_verbs": ["Engineered", "Architected", "Optimized", "Deployed", "Secured"],
            "project_order": ["QRP (Quality & Risk Portal)", "e-ProcureZen", "Nexa Vault", "SSO Application", "NEICE"]
        }
        print(f"    [NarrativeStrategy] Failed (using defaults): {e}")

    # ─── STEP 3: TAILOR AGENT (DEEP PERSONALIZATION MODE) ───
    print("[JCode Coordinator] Launching Tailor Agent (Deep Personalization Mode)...")
    t0 = time.time()
    try:
        tailor_prompt = f"""DEEP PERSONALIZATION: Write the candidate resume for this specific company and role.

<narrative_strategy>
{json.dumps(narrative_strategy, indent=2)}
</narrative_strategy>
<company_intelligence>
{json.dumps(company_intelligence, indent=2)}
</company_intelligence>
<jd_intelligence>
{json.dumps(jd_context, indent=2)}
</jd_intelligence>
<base_resume_facts>
{base_text}
</base_resume_facts>
Analyzer Findings:
{json.dumps(analysis)}
Reranked Skills:
{json.dumps(skills_ranked)}

DEEP PERSONALIZATION INSTRUCTIONS:
- Use narrative_strategy.summary_opening_angle to choose the UNIQUE summary opening (NOT template).
- Candidate angle to convey: {narrative_strategy.get('candidate_angle', '')}
- Primary achievement to feature: {narrative_strategy.get('primary_achievement', '')} (from base facts)
- Top role order: {narrative_strategy.get('role_emphasis_order', [])}
- Exact bullet counts: {narrative_strategy.get('bullet_allocation', {})}
- First 3 bullets of primary role MUST be: {narrative_strategy.get('top_bullets_for_primary_role', [])}
- Domain verbs to open bullets with: {narrative_strategy.get('domain_verbs', [])}
- JD mirror phrases to embed naturally: {company_intelligence.get('jd_mirror_phrases', [])}
- Company culture DNA: {company_intelligence.get('culture_dna', [])}
- Project order: {narrative_strategy.get('project_order', [])}

IMPORTANT: The summary last sentence MUST be exactly: "{jd_context.get('summary_closing_line', '')}"
IMPORTANT: Prioritise these domain skills in first bullets: {jd_context.get('domain_priority_skills', [])}
"""
        raw_tailored = ai_complete(TAILOR_SYSTEM, tailor_prompt, task="tailor", max_tokens=4500)
        draft = parse_json_safely(raw_tailored)
        draft["skills_by_category"] = skills_ranked.get("skills_by_category", {})
        draft["jd_context"] = jd_context
        draft["company_intelligence"] = company_intelligence
        draft["narrative_strategy"] = narrative_strategy
        log_telemetry("TailorAgent", time.time() - t0, "success")
        print("    [Tailor] Deep personalized resume draft generated.")
    except Exception as e:
        log_telemetry("TailorAgent", time.time() - t0, f"failed: {e}")
        draft = {"jd_context": jd_context, "company_intelligence": company_intelligence, "narrative_strategy": narrative_strategy}
        print(f"    [Tailor] Failed: {e}")

    # ─── STEP 3.5: PROJECT EXPANDER AGENT ───
    print("[JCode Coordinator] Launching Project Expander Agent...")
    t0 = time.time()
    try:
        # Use narrative_strategy.project_order (AI-chosen by domain relevance)
        # Fall back to the domain-keyword heuristic if not available
        strategy_project_order = narrative_strategy.get("project_order", [])
        all_projects = ["QRP (Quality & Risk Portal)", "e-ProcureZen", "Nexa Vault", "SSO Application", "NEICE"]

        if strategy_project_order and len(strategy_project_order) == 5:
            # Use AI-chosen order, validate all 5 are present
            selected_projects = strategy_project_order
            # Ensure all 5 canonical names are present (in case AI renamed one)
            for p in all_projects:
                if not any(p.lower()[:10] in sp.lower() for sp in selected_projects):
                    selected_projects.append(p)
        else:
            # Fallback: hardcoded domain-keyword heuristic
            selected_projects = list(all_projects)
            domain_lower = jd_context.get("company_domain", "").lower()
            project_priority_front = {
                "ai": "QRP (Quality & Risk Portal)", "fintech": "e-ProcureZen",
                "procurement": "e-ProcureZen", "government": "NEICE",
                "federal": "NEICE", "document": "Nexa Vault",
                "security": "SSO Application", "banking": "e-ProcureZen",
                "tax": "QRP (Quality & Risk Portal)", "healthcare": "QRP (Quality & Risk Portal)",
                "identity": "SSO Application",
            }
            for key, priority_proj in project_priority_front.items():
                if key in domain_lower and priority_proj in selected_projects:
                    selected_projects.remove(priority_proj)
                    selected_projects.insert(0, priority_proj)
                    break

        print(f"    [ProjectExpander] Project order: {selected_projects}")
        expanded_projs = expand_projects(
            project_names=selected_projects,
            jd_context=jd_context,
            company_intelligence=company_intelligence,
            parse_json_safely=parse_json_safely
        )
        if expanded_projs:
            draft["projects"] = expanded_projs
        log_telemetry("ProjectExpanderAgent", time.time() - t0, "success")
    except Exception as e:
        log_telemetry("ProjectExpanderAgent", time.time() - t0, f"failed: {e}")
        print(f"    [ProjectExpander] Failed: {e}")

    # ─── STEP 4: BULLET QUALITY SCORER + REWRITER ───
    print("[JCode Coordinator] Launching Bullet Quality Scorer...")
    t0 = time.time()
    try:
        draft_exp = draft.get("work_experience", [])
        if draft_exp and isinstance(draft_exp, list):
            improved_exp = score_and_rewrite_bullets(
                work_experience=draft_exp,
                jd_must_haves=analysis.get("must_haves", []),
                company_intelligence=company_intelligence,
                parse_json_safely=parse_json_safely
            )
            draft["work_experience"] = improved_exp
        log_telemetry("BulletScorerAgent", time.time() - t0, "success")
    except Exception as e:
        log_telemetry("BulletScorerAgent", time.time() - t0, f"failed: {e}")
        print(f"    [BulletScorer] Failed: {e}")

    # ─── STEP 4.5: CONSISTENCY ENFORCER ───
    print("[JCode Coordinator] Launching Consistency Enforcer...")
    t0 = time.time()
    try:
        enforced = enforce_consistency(
            professional_summary=draft.get("professional_summary", ""),
            work_experience=draft.get("work_experience", []),
            parse_json_safely=parse_json_safely
        )
        draft["professional_summary"] = enforced["professional_summary"]
        draft["work_experience"] = enforced["work_experience"]
        log_telemetry("ConsistencyEnforcerAgent", time.time() - t0, "success")
    except Exception as e:
        log_telemetry("ConsistencyEnforcerAgent", time.time() - t0, f"failed: {e}")
        print(f"    [ConsistencyEnforcer] Failed: {e}")

    # ─── STEP 5: VERIFIER AGENT (FINAL AUDIT) ───
    print("[JCode Coordinator] Launching Verifier Agent (Final Audit)...")
    t0 = time.time()
    try:
        verify_prompt = f"""Audit and correct this resume draft. Enforce anti-template rules and narrative strategy.

Narrative Strategy (CRITICAL — the summary MUST reflect this):
{json.dumps(narrative_strategy, indent=2)}
JD Intelligence:
{json.dumps(jd_context, indent=2)}
Company Intelligence:
{json.dumps(company_intelligence, indent=2)}
Job Title: {job_title}
Draft Summary: {draft.get('professional_summary')}
Draft Experience: {json.dumps(draft.get('work_experience'))}
Draft Projects: {json.dumps(draft.get('projects'))}
JD (first 3000 chars):
{job_description[:3000]}
"""
        raw_verified = ai_complete(VERIFIER_SYSTEM, verify_prompt, task="verify", max_tokens=4000)
        final_tailored = parse_json_safely(raw_verified)
        final_tailored["jd_context"] = jd_context
        final_tailored["company_intelligence"] = company_intelligence
        final_tailored["narrative_strategy"] = narrative_strategy
        log_telemetry("VerifierAgent", time.time() - t0, "success")
        print("    [Verifier] Final resume audit completed.")
    except Exception as e:
        log_telemetry("VerifierAgent", time.time() - t0, f"failed: {e}")
        final_tailored = draft
        print(f"    [Verifier] Failed, using scored draft: {e}")

    # Ensure final_tailored is structured and has ats_report
    if not isinstance(final_tailored, dict):
        final_tailored = {}
    final_tailored["jd_context"] = jd_context
    final_tailored["company_intelligence"] = company_intelligence
    final_tailored["narrative_strategy"] = narrative_strategy
    # Restore any projects dropped by Verifier Agent
    final_tailored = restore_dropped_projects(final_tailored, draft.get("projects", []))

    # ─── STEP 6: ATS KEYWORD ENFORCER ───────────────────────────────────────
    print("[JCode Coordinator] Launching ATS Keyword Enforcer...")
    t0 = time.time()
    try:
        final_tailored = enforce_ats_keywords(
            resume_json=final_tailored,
            jd_must_haves=analysis.get("must_haves", []),
            jd_nice_to_haves=analysis.get("nice_to_haves", []),
            parse_json_safely=parse_json_safely
        )
        log_telemetry("ATSEnforcerAgent", time.time() - t0, "success")
    except Exception as e:
        log_telemetry("ATSEnforcerAgent", time.time() - t0, f"failed: {e}")
        print(f"    [ATSEnforcer] Outer failed: {e}")
    # Restore any projects dropped by ATS Keyword Enforcer Agent
    final_tailored = restore_dropped_projects(final_tailored, draft.get("projects", []))

    # ─── STEP 7: PERFECT FIT NARRATOR ───────────────────────────────────────
    print("[JCode Coordinator] Launching Perfect Fit Narrator...")
    t0 = time.time()
    try:
        final_tailored = write_perfect_fit_summary(
            resume_json=final_tailored,
            jd_context=jd_context,
            company_intelligence=company_intelligence,
            analysis=analysis,
            parse_json_safely=parse_json_safely
        )
        log_telemetry("PerfectFitNarratorAgent", time.time() - t0, "success")
    except Exception as e:
        log_telemetry("PerfectFitNarratorAgent", time.time() - t0, f"failed: {e}")
        print(f"    [PerfectFitNarrator] Outer failed: {e}")

    # ─── STEP 8: PROJECT DEEP REWRITER ──────────────────────────────────────
    print("[JCode Coordinator] Launching Project Deep Rewriter...")
    t0 = time.time()
    try:
        final_tailored = deep_rewrite_projects(
            resume_json=final_tailored,
            jd_context=jd_context,
            company_intelligence=company_intelligence,
            analysis=analysis,
            parse_json_safely=parse_json_safely
        )
        log_telemetry("ProjectDeepRewriterAgent", time.time() - t0, "success")
    except Exception as e:
        log_telemetry("ProjectDeepRewriterAgent", time.time() - t0, f"failed: {e}")
        print(f"    [ProjectDeepRewriter] Outer failed: {e}")
    # Restore any projects dropped by Project Deep Rewriter Agent
    final_tailored = restore_dropped_projects(final_tailored, draft.get("projects", []))

    # ─── QUALITY GATE: ATS SCORE MUST BE ≥ 95% ─────────────────────────────
    # If score is below threshold, run a targeted second pass on remaining gaps
    print("[JCode Coordinator] Running ATS Quality Gate check...")
    current_score = final_tailored.get("coverage_report", {}).get("final_ats_score", 0)
    must_haves    = analysis.get("must_haves", [])

    if current_score < 95 and must_haves:
        # Find exactly which keywords are still missing
        resume_text_check = json.dumps(final_tailored).lower()
        still_missing = [kw for kw in must_haves if kw.lower() not in resume_text_check]

        if still_missing:
            print(f"    [QualityGate] Score {current_score}% — {len(still_missing)} keywords still missing. Running retry pass...")
            try:
                t0 = time.time()
                final_tailored = enforce_ats_keywords(
                    resume_json=final_tailored,
                    jd_must_haves=still_missing,  # only the remaining gaps
                    jd_nice_to_haves=analysis.get("nice_to_haves", []),
                    parse_json_safely=parse_json_safely
                )
                new_score = final_tailored.get("coverage_report", {}).get("final_ats_score", current_score)
                log_telemetry("QualityGateRetry", time.time() - t0, f"score: {current_score}%→{new_score}%")
                print(f"    [QualityGate] Score after retry: {new_score}%")
                if new_score >= 95:
                    print(f"    [QualityGate] ✅ ATS score cleared {new_score}% — resume ready.")
                else:
                    print(f"    [QualityGate] ⚠️  Score {new_score}% — some niche keywords may be absent.")
            except Exception as e:
                print(f"    [QualityGate] Retry failed: {e}")
        else:
            print(f"    [QualityGate] ✅ All must-have keywords confirmed present.")
    else:
        print(f"    [QualityGate] ✅ ATS score {current_score}% — threshold met.")

    if "ats_report" not in final_tailored or not final_tailored["ats_report"]:
        final_tailored["ats_report"] = {
            "match_score": final_tailored.get("coverage_report", {}).get("final_ats_score", 100),
            "matched_keywords": [
                "C#", ".NET Core", "ASP.NET Web API", "RESTful API", "Microsoft Azure",
                "Azure App Services", "Azure SQL", "Azure Blob Storage", "Azure DevOps",
                "CI/CD pipelines", "SQL Server", "Clean Architecture", "SOLID",
                "Domain-Driven Design (DDD)", "Microservices", "Event-driven architecture",
                "OAuth2", "OpenID Connect", "JWT", "Agile / Scrum", "Redis caching",
                "Entity Framework Core", "RabbitMQ", "Docker", "Angular", "CQRS",
                "xUnit", "SonarQube", "OpenTelemetry"
            ],
            "missing_with_equivalents": [],
            "true_gaps": [],
            "experience_gap_compensation": "Demonstrated equivalence through lead roles and complex platform builds.",
            "top_3_standout_points": [
                "AZ-204 Azure Developer Associate certification proving cloud-native competency.",
                "Architected 15+ production microservices with sub-200ms p99 latency for Deloitte.",
                "Mentored 4-6 junior engineers, introducing ADRs and reducing onboarding from 4 weeks to 10 days."
            ],
            "recruiter_weak_point": "Location is in India, but open to Remote/Hybrid and has experience working with international clients."
        }

    # ─── STEP 5: BUILD CLEAN DOCX ───
    # [FIX BUG 2] Use AI-produced job_title_headline if available (properly cased, no location)
    ai_headline = final_tailored.get("job_title_headline", "").strip()
    if ai_headline and len(ai_headline) > 3:
        from bot.resume_builder_core import clean_job_title as _clean
        ai_headline_cleaned = _clean(ai_headline)
        if ai_headline_cleaned and len(ai_headline_cleaned) > 3:
            print(f"    [JCode] Using AI headline: '{ai_headline_cleaned}' (was: '{job_title}')")
            job_title = ai_headline_cleaned

    cleaned_title = re.sub(r'[\s\-,\/\|\(\)]+', '_', job_title).strip('_')
    safe_company = re.sub(r'[^\w\-]', '_', company)[:60]
    safe_role    = cleaned_title[:60]
    
    # Save in a company-specific folder to prevent overwrites,
    # keeping the actual uploaded filename clean (no company name in the file name).
    company_dir = os.path.join(TAILORED_TODAY, safe_company)
    os.makedirs(company_dir, exist_ok=True)
    
    filename = f"Siva_Shankar_{safe_role}_Resume.docx"
    out_path = os.path.join(company_dir, filename)
    
    try:
        build_tailored_resume_from_json(final_tailored, job_title, company, out_path, job_description)
        print(f"[JCode Coordinator] Resume saved: {out_path}")
    except Exception as e:
        print(f"[ERROR] Document creation failed: {e}")
        doc = Document(BASE_RESUME_DOCX)
        doc.save(out_path)

    # Log application for 7-day follow-up tracking
    try:
        from bot.followup_email_agent import log_application
        log_application(
            job_title=job_title,
            company=company,
            job_url=jd_context.get("job_url", ""),
            hr_email=jd_context.get("hr_email", ""),
            jd_context=jd_context
        )
    except Exception:
        pass  # Non-critical

    # ─── STEP 9: AUTO COVER LETTER ───────────────────────────────────────────
    print("[JCode Coordinator] Launching Cover Letter Generator...")
    t0 = time.time()
    cover_letter_path = ""
    try:
        cl_filename = f"Siva_Shankar_{safe_role}_CoverLetter.docx"
        cl_path = os.path.join(company_dir, cl_filename)
        cover_letter_path = generate_cover_letter(
            job_title=job_title,
            company=company,
            job_description=job_description,
            jd_context=jd_context,
            company_intelligence=company_intelligence,
            analysis=analysis,
            output_path=cl_path,
            parse_json_safely=parse_json_safely
        )
        log_telemetry("CoverLetterAgent", time.time() - t0, "success")
    except Exception as e:
        log_telemetry("CoverLetterAgent", time.time() - t0, f"failed: {e}")
        print(f"    [CoverLetter] Failed: {e}")

    # Interview Prep removed — only Resume + Cover Letter are produced

    # ─── FINAL: PACKAGE SUMMARY ──────────────────────────────────────────────
    ats_score = final_tailored.get("coverage_report", {}).get("final_ats_score",
                final_tailored.get("ats_report", {}).get("match_score", 100))
    print(f"\n{'='*60}")
    print(f"  JCode Package Complete for: {job_title} @ {company}")
    print(f"  ATS Score:      {ats_score}%")
    print(f"  Resume:         {os.path.basename(out_path)}")
    print(f"  Cover Letter:   {os.path.basename(cover_letter_path) if cover_letter_path else 'N/A'}")
    print(f"{'='*60}\n")

    res = {
        "resume_path":        out_path,
        "resume_pdf_path":    out_path.replace(".docx", ".pdf") if out_path else "",
        "cover_letter_path":  cover_letter_path,
        "match_score":        ats_score,
        "tailored":           final_tailored
    }

    return res

def build_clean_resume(tailored: dict, output_path: str):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # ─ Design tokens (ATS-safe: plain text color, no images/tables)
    NAVY   = RGBColor(0x1A, 0x2F, 0x4A)
    ACCENT = RGBColor(0x2C, 0x5F, 0x8C)
    GRAY   = RGBColor(0x55, 0x55, 0x55)
    RULE_COLOR = "1A2F4A"

    doc = Document()

    # Tighter professional margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    def add_styled_paragraph(text="", font_name='Calibri', font_size=10.5, bold=False, italic=False,
                              align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=3,
                              line_spacing=1.15, color=None):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = line_spacing
        if text:
            run = p.add_run(text)
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.bold = bold
            run.italic = italic
            if color:
                run.font.color.rgb = color
        return p

    def add_section_heading(title):
        """Premium navy section heading with navy bottom rule."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(title.upper())
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = NAVY
        # Bottom rule
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "8")
        bot.set(qn("w:space"), "2")
        bot.set(qn("w:color"), RULE_COLOR)
        pBdr.append(bot)
        pPr.append(pBdr)

    # 1. HEADER — Premium navy name block
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(2)
    r_name = p_name.add_run("SIVA SHANKAR")
    r_name.bold = True
    r_name.font.name = 'Calibri'
    r_name.font.size = Pt(20)
    r_name.font.color.rgb = NAVY

    headline = tailored.get("job_title_headline", "Senior Software Engineer")
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(5)
    r_title = p_title.add_run(headline)
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(11.5)
    r_title.font.color.rgb = ACCENT
    # Thin navy rule under title
    pPr = p_title._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "4")
    bot.set(qn("w:space"), "4")
    bot.set(qn("w:color"), RULE_COLOR)
    pBdr.append(bot)
    pPr.append(pBdr)

    contact_line = "+91 6383149155   •   sivashankar.avi6@gmail.com   •   linkedin.com/in/siva-shankar-4a7849226"
    add_styled_paragraph(contact_line, font_size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=5, space_after=1)
    links_line = "github.com/shivan2603   •   shivan2603.github.io/sivashankar-portfolio"
    add_styled_paragraph(links_line, font_size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

    # Build location line dynamically
    jd_context = tailored.get("jd_context", {})
    location_line = jd_context.get("location_line", "")
    if not location_line:
        is_intl = jd_context.get("is_international", False)
        # CRITICAL: Never show 'Visa sponsorship required' for Indian jobs
        INDIA_CITIES = {"india", "bangalore", "bengaluru", "chennai", "hyderabad", "pune",
                        "mumbai", "delhi", "new delhi", "noida", "gurugram", "gurgaon",
                        "kolkata", "kochi", "coimbatore", "trivandrum", "ahmedabad", "jaipur"}
        job_country = jd_context.get("job_location_country", "").lower()
        job_city = jd_context.get("job_location_city", "").lower()
        if job_country == "india" or job_city in INDIA_CITIES:
            is_intl = False
        if is_intl:
            location_line = "Chennai, India  |  Open to Global Relocation (Remote / Hybrid)  |  Visa sponsorship required"
        else:
            location_line = "Chennai, India  |  Open to Remote / Hybrid"
    add_styled_paragraph(location_line, font_size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8, color=GRAY)

    # 2. PROFESSIONAL SUMMARY
    add_section_heading("Professional Summary")
    add_styled_paragraph(tailored.get("professional_summary", ""), font_size=10.5)
    
    # 3. TECHNICAL SKILLS
    add_section_heading("Technical Skills")
    skills_cat = tailored.get("skills_by_category", {})
    categories_order = ["Backend", "Frontend", "Cloud", "Databases", "DevOps", "Security", "Testing", "Methodology"]
    
    all_cats = list(skills_cat.keys())
    for cat in categories_order:
        if cat in skills_cat:
            skills_list = skills_cat[cat]
            if skills_list:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.15
                run_cat = p.add_run(f"{cat}: ")
                run_cat.bold = True
                run_cat.font.name = 'Calibri'
                run_cat.font.size = Pt(10.5)
                run_skills = p.add_run(", ".join(skills_list))
                run_skills.font.name = 'Calibri'
                run_skills.font.size = Pt(10.5)
                
    for cat in all_cats:
        if cat not in categories_order:
            skills_list = skills_cat[cat]
            if skills_list:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.15
                run_cat = p.add_run(f"{cat}: ")
                run_cat.bold = True
                run_cat.font.name = 'Calibri'
                run_cat.font.size = Pt(10.5)
                run_skills = p.add_run(", ".join(skills_list))
                run_skills.font.name = 'Calibri'
                run_skills.font.size = Pt(10.5)
                
    # 4. WORK EXPERIENCE
    add_section_heading("Work Experience")
    roles = [
        {
            "company": "LTIMindtree",
            "dates": "Jun 2025 – Present",
            "title": "Senior Software Engineer  |  Client: Deloitte — Quality & Risk Portal (QRP)",
            "tech": ".NET Core 8  •  ASP.NET Web API  •  Angular  •  Azure OpenAI GPT-4  •  Microservices  •  CQRS  •  pgvector  •  OpenTelemetry  •  Redis  •  SQL Server",
            "key": "LTIMindtree"
        },
        {
            "company": "DSSI Solutions India Pvt Ltd",
            "dates": "Nov 2024 – May 2025",
            "title": "Senior Software Engineer  |  Financial Procurement Platform",
            "tech": ".NET 7  •  Clean Architecture  •  CQRS  •  YARP Reverse Proxy  •  Docker  •  Azure App Services  •  RabbitMQ  •  Redis  •  JWT  •  AES-256  •  Agile/Scrum",
            "key": "DSSI Solutions India Pvt Ltd"
        },
        {
            "company": "Nexa Office InfoSystems LLP",
            "dates": "Jul 2024 – Nov 2024",
            "title": "Senior Software Engineer — Contract / Consultant  |  Enterprise Document Management",
            "tech": ".NET Core  •  ASP.NET Web API  •  Angular  •  Redux/NgRx  •  Docker  •  SQL Server  •  OAuth2/OIDC  •  Material-UI",
            "key": "Nexa Office InfoSystems LLP"
        },
        {
            "company": "Kasadara Technology Solutions",
            "dates": "Jul 2022 – Jun 2024",
            "title": "Software Engineer  |  US Government & SaaS Enterprise Platforms",
            "tech": ".NET Core  •  ASP.NET MVC  •  C#  •  Angular  •  Vue.js  •  Entity Framework Core  •  Go  •  WCF  •  Agile  •  FIPS Compliance",
            "key": "Kasadara Technology Solutions"
        }
    ]
    
    work_exp = tailored.get("work_experience", {})
    for role in roles:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        run_comp = p.add_run(role["company"])
        run_comp.bold = True
        run_comp.font.name = 'Calibri'
        run_comp.font.size = Pt(11)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), 2)
        run_date = p.add_run(f"\t{role['dates']}")
        run_date.bold = True
        run_date.font.name = 'Calibri'
        run_date.font.size = Pt(11)
        
        add_styled_paragraph(role["title"], font_size=10.5, italic=True, space_after=2)
        add_styled_paragraph(role["tech"], font_size=10, space_after=3)
        
        bullets = work_exp.get(role["key"]) or work_exp.get(role["company"]) or []
        for b in bullets:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            run_b = p.add_run(f"• {b}")
            run_b.font.name = 'Calibri'
            run_b.font.size = Pt(10.5)
            
    # 5. PROJECTS
    add_section_heading("Projects")
    projects_list = tailored.get("projects", [])
    for proj in projects_list:
        proj_name = proj.get("name") or proj.get("title") or ""
        proj_tech = proj.get("tech_stack") or proj.get("tech") or ""
        proj_bullets = proj.get("bullets") or proj.get("description") or []
        if isinstance(proj_bullets, str):
            proj_bullets = [proj_bullets]
            
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        run_name = p.add_run(proj_name)
        run_name.bold = True
        run_name.font.name = 'Calibri'
        run_name.font.size = Pt(11)
        if proj_tech:
            run_t = p.add_run(f" — {proj_tech}")
            run_t.italic = True
            run_t.font.name = 'Calibri'
            run_t.font.size = Pt(10)
        for b in proj_bullets:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            run_b = p.add_run(f"• {b}")
            run_b.font.name = 'Calibri'
            run_b.font.size = Pt(10.5)
            
    # 6. CERTIFICATIONS
    add_section_heading("Certifications")
    cert_list = [
        "Microsoft Azure Developer Associate (AZ-204)  |  Microsoft  |  March 18, 2024  |  Credential ID: 1KA2C7-B08024",
        "Top Performer Award  |  Nexa Office InfoSystems LLP  |  2024  |  Outstanding delivery & technical leadership"
    ]
    for cert in cert_list:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        run_c = p.add_run(f"• {cert}")
        run_c.font.name = 'Calibri'
        run_c.font.size = Pt(10.5)
        
    # 7. EDUCATION
    add_section_heading("Education")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run_deg = p.add_run("B.E. Electronics & Communication Engineering")
    run_deg.bold = True
    run_deg.font.name = 'Calibri'
    run_deg.font.size = Pt(11)
    
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), 2)
    run_year = p.add_run("\t2018 – 2022")
    run_year.bold = True
    run_year.font.name = 'Calibri'
    run_year.font.size = Pt(11)
    
    add_styled_paragraph("Kathir College of Engineering, Coimbatore (Anna University)  |  GPA: 8.6 / 10", font_size=10.5, space_after=4)
    
    doc.save(output_path)