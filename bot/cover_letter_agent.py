"""
cover_letter_agent.py — Auto Cover Letter Generator
═══════════════════════════════════════════════════
Generates a personalized, compelling cover letter DOCX for every job application.
Uses the full JD Intelligence + Company Intelligence to write a letter that:
  - Opens with a hook referencing the company's specific challenge
  - Stacks 2 proof paragraphs with metrics directly matching JD requirements
  - Closes with urgency and a clear call to action
  - Feels hand-written, not templated
"""

import os, json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bot.ai_router import ai_complete

# ─── DESIGN TOKENS (matching resume design) ──────────────────────────────────
NAVY   = RGBColor(0x1A, 0x2F, 0x4A)
ACCENT = RGBColor(0x2C, 0x5F, 0x8C)
GRAY   = RGBColor(0x55, 0x55, 0x55)

COVER_LETTER_SYSTEM = """You are the Cover Letter Writer Agent — you write personalized, compelling,
recruiter-stopping cover letters that make hiring managers WANT to interview the candidate immediately.

YOUR COVER LETTER MUST:
1. Feel hand-written and genuine — NOT like a template
2. Reference the specific company and role by name
3. Open with a hook that grabs attention in the first sentence
4. Prove fit with 2-3 specific metrics from the candidate's experience
5. Mirror the JD's exact language and values
6. Close with confident urgency — not begging, but assertive

COVER LETTER STRUCTURE (4 paragraphs, ~250-300 words total):

PARAGRAPH 1 — THE HOOK (2-3 sentences):
- Don't open with "I am writing to apply for..." — that's instant trash can.
- Open with an insight about the company's challenge OR a bold achievement statement.
- Examples:
  * "When [Company] announced [product/initiative], it confirmed what I've been building toward for 4 years..."
  * "Twelve production microservices. 99.98% uptime. Zero incidents in 6 months. That's the standard I hold myself to — and exactly what [Company] needs for [role]."
- End the paragraph by stating why THIS company specifically.

PARAGRAPH 2 — PROOF BLOCK 1 (3-4 sentences):
- Address the JD's #1 technical requirement with a specific story + metric.
- Format: Context → Action → Result (CAR framework)
- Include an exact number. No vague claims.
- Mirror the JD's exact phrasing where possible.

PARAGRAPH 3 — PROOF BLOCK 2 (3-4 sentences):
- Address the JD's #2 requirement OR show cultural fit.
- If team lead role: include mentoring achievement.
- If individual contributor: include depth/architecture decision.
- Show you understand the COMPANY's domain, not just the technology.

PARAGRAPH 4 — CLOSING CTA (2-3 sentences):
- State availability and notice period (serving notice, LWD: 14th August 2026).
- Express genuine excitement — but NOT desperation.
- End with a confident CTA: "I would welcome the opportunity to discuss..."
- Sign off: "Yours sincerely, Siva Shankar"

TONE RULES:
- Confident but not arrogant
- Specific and evidence-based — never vague
- First-person ("I") is fine in cover letters
- No buzzwords: "synergy", "passionate", "go-getter", "team player" — show don't tell

CANDIDATE FACTS (use only these — never invent):
- Name: Siva Shankar V
- Phone: +91 6383149155
- Email: sivashankar.avi6@gmail.com
- Current: Senior Software Engineer at LTIMindtree (Client: Deloitte Enterprise Tax Platform)
- Stack: C#, .NET Core 8, ASP.NET Web API, Azure, Microservices, CQRS, Angular, Redis, SQL Server
- AZ-204 certified
- 4+ years experience
- Notice period: Serving notice, Last Working Day: 14th August 2026
- Key metrics: sub-100ms p99 latency, 12+ microservices, 99.98% uptime, 38% query latency reduction,
  60% manual effort saved, 45% DB load reduction, 85% test coverage, 3x throughput

Return ONLY valid JSON:
{
  "greeting": "Dear Hiring Manager,",
  "paragraph_1": "...",
  "paragraph_2": "...",
  "paragraph_3": "...",
  "paragraph_4": "...",
  "sign_off": "Yours sincerely,\\nSiva Shankar V\\n+91 6383149155 | sivashankar.avi6@gmail.com"
}"""


def generate_cover_letter(
    job_title: str,
    company: str,
    job_description: str,
    jd_context: dict,
    company_intelligence: dict,
    analysis: dict,
    output_path: str,
    parse_json_safely=None
) -> str:
    """
    Generate a personalized cover letter DOCX for this specific job application.
    Returns the path to the saved DOCX file.
    """
    if parse_json_safely is None:
        def parse_json_safely(raw):
            import json, re
            raw = raw.strip()
            m = re.search(r'(\{.*\})', raw, re.DOTALL)
            candidate = m.group(1).strip() if m else raw
            try:
                return json.loads(candidate, strict=False)
            except Exception:
                return {}

    print(f"    [CoverLetter] Writing personalized cover letter for {company}...")

    try:
        prompt = (
            f"Job Title: {job_title}\n"
            f"Company: {company}\n"
            f"Company Domain: {jd_context.get('company_domain', 'technology')}\n"
            f"JD Must-Haves: {analysis.get('must_haves', [])}\n"
            f"JD Mirror Phrases: {company_intelligence.get('jd_mirror_phrases', [])}\n"
            f"Company Culture DNA: {company_intelligence.get('culture_dna', [])}\n"
            f"Top Differentiators: {company_intelligence.get('top_3_differentiators', [])}\n"
            f"Resume Narrative Angle: {company_intelligence.get('resume_narrative_angle', '')}\n"
            f"Is Team Lead Role: {jd_context.get('is_team_lead_role', False)}\n"
            f"Is International: {jd_context.get('is_international', False)}\n"
            f"Location: {jd_context.get('job_location_city', 'Unknown')}, {jd_context.get('job_location_country', '')}\n\n"
            f"JD Excerpt:\n{job_description[:2000]}\n\n"
            f"Write a compelling 4-paragraph cover letter following the exact structure. "
            f"Make it feel like Siva personally researched {company} and wrote this specifically for them."
        )
        raw = ai_complete(COVER_LETTER_SYSTEM, prompt, task="tailor", max_tokens=1200)
        letter = parse_json_safely(raw)
    except Exception as e:
        print(f"    [CoverLetter] AI generation failed: {e}. Using structured fallback.")
        domain = jd_context.get("company_domain", "technology")
        letter = {
            "greeting": "Dear Hiring Manager,",
            "paragraph_1": (
                f"When I reviewed the {job_title} role at {company}, one thing was immediately clear: "
                f"the challenges your team is solving — {domain} systems at scale — are exactly what I've "
                f"spent the last 4 years building expertise in. At LTIMindtree, delivering for Deloitte's "
                f"enterprise tax platform, I learned what it takes to ship production software that never fails."
            ),
            "paragraph_2": (
                f"Most recently, I profiled and refactored 30+ ASP.NET Web API endpoints to eliminate N+1 "
                f"query loops, achieving sub-100ms p99 latency under peak enterprise load — validated end-to-end "
                f"with OpenTelemetry distributed tracing. I also deployed an Azure OpenAI GPT-4 embedding pipeline "
                f"that automated 60% of previously manual schema migration effort. These aren't side projects — "
                f"they ran in production for Deloitte."
            ),
            "paragraph_3": (
                f"At DSSI Solutions, I architected 12+ procurement microservices using .NET 7, CQRS, and Clean "
                f"Architecture, achieving 99.98% uptime SLA. I also mentored 4 junior engineers on Git Flow and "
                f"CQRS patterns, reducing post-deployment defects by 40%. I don't just write code — I raise the "
                f"engineering standard of the team around me."
            ),
            "paragraph_4": (
                f"I'm currently serving my notice period with a last working day of 14th August 2026, and I'm "
                f"actively looking for my next challenge. I would welcome the opportunity to discuss how my "
                f"background in cloud-native .NET development and enterprise delivery aligns with {company}'s goals. "
                f"Thank you for your time and consideration."
            ),
            "sign_off": "Yours sincerely,\nSiva Shankar V\n+91 6383149155 | sivashankar.avi6@gmail.com"
        }

    # Build the DOCX
    try:
        _build_cover_letter_docx(letter, job_title, company, output_path)
        print(f"    [CoverLetter] Saved: {output_path}")
        return output_path
    except Exception as e:
        print(f"    [CoverLetter] DOCX build failed: {e}")
        return ""


def _build_cover_letter_docx(letter: dict, job_title: str, company: str, output_path: str):
    """Build a premium-looking cover letter DOCX matching the resume design."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    def para(text, bold=False, italic=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT,
             sb=0, sa=6, ls=1.15, color=None):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(sb)
        p.paragraph_format.space_after  = Pt(sa)
        p.paragraph_format.line_spacing  = ls
        if text:
            r = p.add_run(text)
            r.bold   = bold
            r.italic = italic
            r.font.name = "Calibri"
            r.font.size = Pt(size)
            if color:
                r.font.color.rgb = color
        return p

    # ── HEADER: Name + contact ──
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_name.paragraph_format.space_after = Pt(2)
    r_name = p_name.add_run("SIVA SHANKAR V")
    r_name.bold = True
    r_name.font.name = "Calibri"
    r_name.font.size = Pt(16)
    r_name.font.color.rgb = NAVY

    para("+91 6383149155  •  sivashankar.avi6@gmail.com  •  linkedin.com/in/siva-shankar-4a7849226",
         size=9.5, color=GRAY, sa=2)
    para("github.com/shivan2603  •  shivan2603.github.io/sivashankar-portfolio",
         size=9.5, color=GRAY, sa=0)

    # Thin navy rule
    p_rule = doc.add_paragraph()
    p_rule.paragraph_format.space_before = Pt(6)
    p_rule.paragraph_format.space_after  = Pt(12)
    pPr = p_rule._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1A2F4A")
    pBdr.append(bot)
    pPr.append(pBdr)

    # Date
    import time
    date_str = time.strftime("%d %B %Y")
    para(date_str, size=10.5, color=GRAY, sa=12)

    # Role + Company
    para(f"Re: Application for {job_title}", bold=True, size=11, color=NAVY, sa=2)
    para(company, size=10.5, color=GRAY, sa=14)

    # Greeting
    para(letter.get("greeting", "Dear Hiring Manager,"), size=11, sa=12)

    # Body paragraphs
    for key in ["paragraph_1", "paragraph_2", "paragraph_3", "paragraph_4"]:
        text = letter.get(key, "")
        if text:
            para(text, size=11, sa=12, ls=1.2)

    # Sign-off
    sign_off_lines = letter.get("sign_off", "Yours sincerely,\nSiva Shankar V").split("\n")
    para(sign_off_lines[0], size=11, sa=4)
    para("", size=11, sa=4)  # gap for signature
    for line in sign_off_lines[1:]:
        para(line, size=10.5, color=GRAY if "@" in line or "+" in line else None, sa=2)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
